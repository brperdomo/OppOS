"""Extract text from RFP attachments (PDF via Nutrient DWS, DOCX locally)."""

from __future__ import annotations

import json
import logging
import os
from pathlib import Path

import httpx

logger = logging.getLogger(__name__)

NUTRIENT_API_URL = "https://api.nutrient.io/build"
MAX_TEXT_PER_FILE = 15_000
MAX_TOTAL_TEXT = 30_000
SCANNABLE_EXTENSIONS = {".pdf", ".docx"}


def _get_api_key() -> str:
    return os.environ.get("NUTRIENT_API_KEY", "")


def extract_text_from_docx(file_path: Path) -> dict:
    """Extract text from a DOCX file locally (no API credits used)."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n".join(paragraphs)
        logger.info("Extracted %d chars from %s (local DOCX parse)", len(full_text), file_path.name)
        return {
            "text": full_text[:MAX_TEXT_PER_FILE],
            "pages": 0,
            "chars": min(len(full_text), MAX_TEXT_PER_FILE),
            "credits_remaining": "n/a (local)",
            "error": None,
        }
    except Exception as e:
        logger.warning("DOCX extraction failed for %s: %s", file_path.name, e)
        return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "n/a", "error": str(e)}


def _call_nutrient_api(file_path: Path, api_key: str, use_ocr: bool = True) -> httpx.Response:
    """Send a PDF to Nutrient DWS API for text extraction."""
    output_config = {
        "type": "json-content",
        "plainText": True,
    }
    if use_ocr:
        output_config["language"] = "english"

    instructions = json.dumps({
        "parts": [{"file": "file"}],
        "output": output_config,
    })

    with open(file_path, "rb") as f:
        return httpx.post(
            NUTRIENT_API_URL,
            headers={"Authorization": f"Bearer {api_key}"},
            data={"instructions": instructions},
            files={"file": (file_path.name, f, "application/pdf")},
            timeout=300.0,
        )


def _parse_nutrient_response(resp: httpx.Response, file_path: Path) -> dict:
    """Parse a successful Nutrient API response into our result dict."""
    data = resp.json()
    pages = data.get("pages", [])
    text_parts = [p.get("plainText", "") for p in pages if p.get("plainText")]
    full_text = "\n\n".join(text_parts)
    remaining = resp.headers.get("x-pspdfkit-remaining-credits", "?")

    logger.info(
        "Extracted %d chars from %s (%d pages) — %s credits remaining",
        len(full_text), file_path.name, len(pages), remaining,
    )
    return {
        "text": full_text[:MAX_TEXT_PER_FILE],
        "pages": len(pages),
        "chars": min(len(full_text), MAX_TEXT_PER_FILE),
        "credits_remaining": remaining,
        "error": None,
    }


def extract_text_from_pdf(file_path: Path) -> dict:
    """Upload a PDF to Nutrient DWS and return extraction result.

    Tries with OCR first. On timeout (408), retries without OCR (plain text
    extraction only — faster, works for born-digital PDFs like bid solicitations).

    Returns dict with keys: text, pages, chars, credits_remaining, error.
    """
    api_key = _get_api_key()
    if not api_key:
        logger.warning("NUTRIENT_API_KEY not set — skipping text extraction")
        return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "?", "error": "No API key"}

    try:
        resp = _call_nutrient_api(file_path, api_key, use_ocr=True)
        resp.raise_for_status()
        return _parse_nutrient_response(resp, file_path)

    except httpx.HTTPStatusError as e:
        if e.response.status_code in (408, 504):
            logger.info("OCR timed out for %s — retrying without OCR...", file_path.name)
            try:
                resp = _call_nutrient_api(file_path, api_key, use_ocr=False)
                resp.raise_for_status()
                result = _parse_nutrient_response(resp, file_path)
                result["error"] = None
                return result
            except httpx.HTTPStatusError as e2:
                msg = f"HTTP {e2.response.status_code} (retry without OCR)"
                logger.warning("Nutrient API retry failed for %s: %s", file_path.name, msg)
                return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "?", "error": msg}
            except Exception as e2:
                logger.warning("Retry failed for %s: %s", file_path.name, e2)
                return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "?", "error": str(e2)}

        msg = f"HTTP {e.response.status_code}"
        logger.warning("Nutrient API error for %s: %s %s", file_path.name, e.response.status_code, e.response.text[:200])
        return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "?", "error": msg}

    except Exception as e:
        logger.warning("Text extraction failed for %s: %s", file_path.name, e)
        return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "?", "error": str(e)}


def extract_file(file_path: Path) -> dict:
    """Route to the right extractor based on file extension."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    return {"text": "", "pages": 0, "chars": 0, "credits_remaining": "n/a", "error": f"Unsupported: {ext}"}


def extract_text_from_attachments(file_paths: list[Path], on_progress=None) -> str:
    """Extract and concatenate text from multiple attachment files.

    on_progress(filename, index, total, result_dict) is called after each file.
    """
    if not file_paths:
        return ""

    all_text = []
    total_len = 0

    pdf_paths = [p for p in file_paths if p.suffix.lower() == ".pdf"]
    if not pdf_paths:
        logger.info("No PDF attachments to extract text from")
        return ""

    for i, path in enumerate(pdf_paths):
        if total_len >= MAX_TOTAL_TEXT:
            logger.info("Reached text limit (%d chars) — skipping remaining files", MAX_TOTAL_TEXT)
            if on_progress:
                on_progress(path.name, i, len(pdf_paths), {"skipped": True})
            break

        result = extract_text_from_pdf(path)
        if on_progress:
            on_progress(path.name, i, len(pdf_paths), result)

        if result["text"]:
            header = f"--- {path.name} ---"
            all_text.append(f"{header}\n{result['text']}")
            total_len += result["chars"]

    return "\n\n".join(all_text)[:MAX_TOTAL_TEXT]
